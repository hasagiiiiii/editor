from docx import Document
from docx.oxml.shared import OxmlElement
from docx.text.paragraph import Paragraph

doc = Document()
doc.add_paragraph('Test Paragraph')

body = doc.element.body
anchor = body[0]

new_p = OxmlElement('w:p')
body.insert(body.index(anchor) + 1, new_p)

p = Paragraph(new_p, doc)
run = p.add_run('World')
run.bold = True

doc.save('d:/editor/scratch_out.docx')
print('Success')
