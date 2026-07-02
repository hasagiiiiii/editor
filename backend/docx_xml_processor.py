"""
docx_xml_processor.py

Module xử lý file .docx bằng can thiệp trực tiếp vào cấu trúc XML.
Sử dụng python-docx cho I/O file và lxml cho thao tác XML,
phục vụ cho ứng dụng chỉnh sửa văn bản bằng AI.

Cách dùng:
    processor = DocxXmlProcessor()

    # Bước 1: Trích xuất đoạn văn kèm ID
    paragraphs = processor.extract_paragraphs_with_ids("input.docx")
    # => {"p_1": "Nội dung đoạn 1", "p_2": "Nội dung đoạn 2", ...}

    # Bước 2: Gửi cho AI chỉnh sửa, nhận lại dict đã sửa
    edited = {"p_1": "Nội dung mới đoạn 1", "p_2": "Nội dung mới đoạn 2"}

    # Bước 3: Gộp nội dung đã sửa vào file gốc
    processor.merge_edited_text("input.docx", edited, "output.docx")

Dependencies:
    pip install python-docx lxml
"""

from __future__ import annotations

import copy
import logging
import shutil
from pathlib import Path
from typing import Any, Dict, List, Tuple

from docx import Document
from lxml import etree

__all__ = ["DocxXmlProcessor"]

logger = logging.getLogger(__name__)


class DocxXmlProcessor:
    """
    Xử lý file .docx bằng can thiệp trực tiếp vào cấu trúc XML.

    Cung cấp hai phương thức chính:
      - extract_paragraphs_with_ids: Trích xuất nội dung từng đoạn văn kèm ID.
      - merge_edited_text: Ghi đè nội dung đã chỉnh sửa vào file gốc,
        giữ nguyên định dạng đoạn (w:pPr) và định dạng ký tự (w:rPr).
    """

    # ── OOXML Namespaces ────────────────────────────────────────────
    _W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    _W = f"{{{_W_NS}}}"
    _XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"

    # ================================================================
    #  INTERNAL HELPERS
    # ================================================================

    def _is_toc_paragraph(self, p_elem: etree._Element) -> bool:
        """
        Kiểm tra xem đoạn văn có thuộc cấu trúc Mục lục (TOC) hay không.

        Nhận diện qua 5 dấu hiệu:
          1. Paragraph style bắt đầu bằng "TOC" (TOC1, TOC2, TOCHeading, ...).
          2. Nằm trong w:sdt có docPartGallery = "Table of Contents".
          3. Chứa w:instrText có chuỗi "TOC".
          4. Chứa w:fldSimple con có w:instr chứa "TOC".
          5. Phần tử cha là w:fldSimple với w:instr chứa "TOC".
        """
        W = self._W

        # 1. Kiểm tra paragraph style
        p_pr = p_elem.find(f"{W}pPr")
        if p_pr is not None:
            p_style = p_pr.find(f"{W}pStyle")
            if p_style is not None:
                style_val = p_style.get(f"{W}val", "")
                if style_val.upper().startswith("TOC"):
                    return True

        # 2. Kiểm tra tổ tiên w:sdt (Structured Document Tag) cho TOC
        ancestor = p_elem.getparent()
        while ancestor is not None:
            if ancestor.tag == f"{W}sdt":
                sdt_pr = ancestor.find(f"{W}sdtPr")
                if sdt_pr is not None:
                    doc_part = sdt_pr.find(f"{W}docPartObj")
                    if doc_part is not None:
                        gallery = doc_part.find(f"{W}docPartGallery")
                        if gallery is not None:
                            val = gallery.get(f"{W}val", "")
                            if "Table of Contents" in val:
                                return True
            ancestor = ancestor.getparent()

        # 3. Kiểm tra w:instrText chứa "TOC" bên trong đoạn
        for instr_text in p_elem.iter(f"{W}instrText"):
            if instr_text.text and "TOC" in instr_text.text.upper():
                return True

        # 4. Kiểm tra w:fldSimple con bên trong đoạn
        for fld in p_elem.iter(f"{W}fldSimple"):
            instr = fld.get(f"{W}instr", "")
            if "TOC" in instr.upper():
                return True

        # 5. Kiểm tra phần tử cha trực tiếp là w:fldSimple
        parent = p_elem.getparent()
        if parent is not None and parent.tag == f"{W}fldSimple":
            instr = parent.get(f"{W}instr", "")
            if "TOC" in instr.upper():
                return True

        return False

    def _extract_paragraph_text(self, p_elem: etree._Element) -> str:
        """
        Trích xuất toàn bộ text thuần túy từ các thẻ w:t bên trong đoạn.

        Ghép nối text từ tất cả w:t (kể cả bên trong w:hyperlink, w:ins, ...)
        theo đúng thứ tự xuất hiện trong XML.
        """
        W = self._W
        fragments: List[str] = []
        for t_elem in p_elem.iter(f"{W}t"):
            if t_elem.text:
                fragments.append(t_elem.text)
        return "".join(fragments)


    def _get_content_paragraphs(
        self, body: etree._Element
    ) -> List[Tuple[etree._Element, str]]:
        """
        Thu thập tất cả đoạn văn có nội dung (không rỗng hoặc có chứa hình ảnh, không thuộc TOC)
        theo đúng thứ tự depth-first recursive walk.
        """
        W = self._W
        results: List[Tuple[etree._Element, str]] = []
        
        def walk(elem: etree._Element):
            if elem.tag == f"{W}p":
                if self._is_toc_paragraph(elem):
                    return
                text = self._extract_paragraph_text(elem)
                has_media = any("drawing" in sub.tag or "shape" in sub.tag or "imagedata" in sub.tag for sub in elem.iter())
                if not text.strip() and not has_media:
                    return
                results.append((elem, text))
            else:
                for child in elem:
                    walk(child)
                    
        walk(body)
        return results

    def _parse_p_pr(self, p_elem: etree._Element) -> Dict[str, str]:
        """
        Phân tích w:pPr để lấy các thuộc tính căn lề, thụt lề và danh sách từ Word XML.
        Trả về dict các style CSS inline tương ứng.
        Nếu đoạn văn thuộc danh sách (w:numPr), thêm key đặc biệt:
          "_list_level": cấp thụt danh sách (0, 1, 2, ...)
        """
        W = self._W
        styles: Dict[str, str] = {}
        
        p_pr = p_elem.find(f"{W}pPr")
        if p_pr is not None:
            # 1. Căn lề (w:jc)
            jc = p_pr.find(f"{W}jc")
            if jc is not None:
                val = jc.get(f"{W}val", "")
                if val == "center":
                    styles["text-align"] = "center"
                elif val == "right":
                    styles["text-align"] = "right"
                elif val == "both":
                    styles["text-align"] = "justify"
                elif val == "left":
                    styles["text-align"] = "left"
            
            # 2. Thụt lề (w:ind)
            ind = p_pr.find(f"{W}ind")
            has_explicit_left = False
            if ind is not None:
                # Lùi dòng trái (w:left)
                left = ind.get(f"{W}left")
                if left:
                    try:
                        px = int(left) / 15
                        styles["margin-left"] = f"{px:.1f}px"
                        has_explicit_left = True
                    except ValueError:
                        pass
                
                # Thụt dòng đầu tiên (w:firstLine)
                first_line = ind.get(f"{W}firstLine")
                if first_line:
                    try:
                        px = int(first_line) / 15
                        styles["text-indent"] = f"{px:.1f}px"
                    except ValueError:
                        pass
                
                # Thụt dòng treo (w:hanging) - dùng cho bullet/number lists
                hanging = ind.get(f"{W}hanging")
                if hanging:
                    try:
                        px = int(hanging) / 15
                        styles["text-indent"] = f"-{px:.1f}px"
                        styles["padding-left"] = f"{px:.1f}px"
                    except ValueError:
                        pass
            
            # 3. Danh sách (w:numPr) - bullet và numbered lists
            num_pr = p_pr.find(f"{W}numPr")
            if num_pr is not None:
                ilvl_elem = num_pr.find(f"{W}ilvl")
                ilvl = 0
                if ilvl_elem is not None:
                    try:
                        ilvl = int(ilvl_elem.get(f"{W}val", "0"))
                    except ValueError:
                        pass
                styles["_list_level"] = str(ilvl)
                
                # Nếu không có explicit left indent từ w:ind, tính theo cấp danh sách
                if not has_explicit_left:
                    indent_px = (ilvl + 1) * 28  # ~28px mỗi cấp
                    styles["margin-left"] = f"{indent_px}px"

            # 4. Khoảng cách và giãn dòng (w:spacing)
            spacing = p_pr.find(f"{W}spacing")
            if spacing is not None:
                before = spacing.get(f"{W}before")
                after = spacing.get(f"{W}after")
                line = spacing.get(f"{W}line")
                
                if before:
                    try:
                        px = int(before) / 15
                        styles["margin-top"] = f"{px:.1f}px"
                    except ValueError:
                        pass
                if after:
                    try:
                        px = int(after) / 15
                        styles["margin-bottom"] = f"{px:.1f}px"
                    except ValueError:
                        pass
                else:
                    styles["margin-bottom"] = "8px"

                if line:
                    try:
                        line_val = int(line)
                        line_rule = spacing.get(f"{W}lineRule", "auto")
                        if line_rule == "auto":
                            styles["line-height"] = f"{line_val / 240:.2f}"
                        else:
                            styles["line-height"] = f"{line_val / 15:.1f}px"
                    except ValueError:
                        pass
            else:
                styles["margin-top"] = "0px"
                styles["margin-bottom"] = "8px"
                styles["line-height"] = "1.15"
                         
        return styles

    def _load_image_relationships(self, docx_path: str) -> Dict[str, str]:
        """
        Đọc và ánh xạ các relationship ID (rId) sang tên file ảnh thực tế.
        """
        import zipfile
        rels: Dict[str, str] = {}
        try:
            with zipfile.ZipFile(docx_path) as z:
                if "word/_rels/document.xml.rels" in z.namelist():
                    rels_data = z.read("word/_rels/document.xml.rels")
                    root = etree.fromstring(rels_data)
                    R_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
                    for rel in root.findall(f"{{{R_NS}}}Relationship"):
                        r_id = rel.get("Id")
                        target = rel.get("Target")
                        rel_type = rel.get("Type", "")
                        if "image" in rel_type.lower() and r_id and target:
                            filename = Path(target).name
                            rels[r_id] = filename
        except Exception as e:
            logger.warning("Không thể đọc quan hệ hình ảnh: %s", e)
        return rels

    def _body_to_html_and_paragraphs(self, body_elem: etree._Element, image_rels: Dict[str, str], doc_filename: str) -> Tuple[str, Dict[str, Dict[str, str]]]:
        """
        Convert body XML to HTML and also return the dictionary of paragraph text and inner HTML.
        This ensures 100% alignment between layout HTML and the paragraph index.
        """
        W = self._W
        paragraphs_dict: Dict[str, Dict[str, str]] = {}
        p_index_ref = [0]
        
        def walk(elem: etree._Element) -> str:
            if elem.tag == f"{W}p":
                if self._is_toc_paragraph(elem):
                    return ""
                text = self._extract_paragraph_text(elem)
                has_media = any("drawing" in sub.tag or "shape" in sub.tag or "imagedata" in sub.tag for sub in elem.iter())
                if not text.strip() and not has_media:
                    return ""
                
                p_index_ref[0] += 1
                p_idx = p_index_ref[0]
                p_id = f"p_{p_idx}"
                
                # Styles and inner HTML
                styles_dict = self._parse_p_pr(elem)
                
                html_parts: List[str] = []
                for child in elem:
                    if child.tag == f"{W}r":
                        html_parts.append(self._run_to_html(child, image_rels, doc_filename))
                    elif child.tag == f"{W}hyperlink":
                        link_text = []
                        for sub_child in child:
                            if sub_child.tag == f"{W}r":
                                link_text.append(self._run_to_html(sub_child, image_rels, doc_filename))
                        html_parts.append("".join(link_text))
                    elif child.tag == f"{W}sdt":
                        sdt_content = child.find(f"{W}sdtContent")
                        if sdt_content is not None:
                            for sub in sdt_content:
                                if sub.tag == f"{W}r":
                                    html_parts.append(self._run_to_html(sub, image_rels, doc_filename))
                
                inner_html = "".join(html_parts)
                paragraphs_dict[p_id] = {
                    "text": text,
                    "html": inner_html or text or "&nbsp;"
                }
                
                # Lọc bỏ key nội bộ _list_level khỏi CSS string
                list_level = styles_dict.pop("_list_level", None)
                styles_str = "; ".join(f"{k}: {v}" for k, v in styles_dict.items())
                
                # Thêm ký tự bullet cho các mục danh sách
                display_html = inner_html or "&nbsp;"
                if list_level is not None:
                    lvl = int(list_level)
                    bullets = ["•", "◦", "▪", "‣", "–"]
                    marker = bullets[lvl % len(bullets)]
                    display_html = f'<span style="margin-right:6px;">{marker}</span>{inner_html}'
                    
                styles_attr = f' style="{styles_str}"' if styles_str else ''
                return f'<div class="para-block" id="block-{p_id}"{styles_attr}>{display_html}</div>'
                
            elif elem.tag == f"{W}tbl":
                tbl_parts = ["<table class=\"doc-table\">"]
                for row in elem.findall(f"{W}tr"):
                    tbl_parts.append("<tr>")
                    for cell in row.findall(f"{W}tc"):
                        tbl_parts.append("<td>")
                        cell_contents = []
                        for child in cell:
                            cell_contents.append(walk(child))
                        tbl_parts.append("".join(cell_contents))
                        tbl_parts.append("</td>")
                    tbl_parts.append("</tr>")
                tbl_parts.append("</table>")
                return "".join(tbl_parts)
                
            else:
                parts = []
                for child in elem:
                    parts.append(walk(child))
                return "".join(parts)

        html_out = walk(body_elem)
        return html_out, paragraphs_dict

    def _run_to_html(self, r_elem: etree._Element, image_rels: Dict[str, str], doc_filename: str) -> str:
        """
        Chuyển đổi một thẻ w:r sang HTML, bảo toàn format bold, italic, u và kích thước ảnh.
        """
        W = self._W
        A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
        R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        V_NS = "urn:schemas-microsoft-com:vml"
        WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
        
        r_pr = r_elem.find(f"{W}rPr")
        is_bold = False
        is_italic = False
        is_underline = False
        
        if r_pr is not None:
            if r_pr.find(f"{W}b") is not None:
                is_bold = True
            if r_pr.find(f"{W}i") is not None:
                is_italic = True
            if r_pr.find(f"{W}u") is not None:
                is_underline = True
                
        run_parts: List[str] = []
        
        for child in r_elem:
            if child.tag == f"{W}t":
                if child.text:
                    run_parts.append(child.text)
            elif child.tag == f"{W}drawing":
                # Trích xuất kích thước ảnh chuẩn từ extent
                extent = child.find(f".//{{{WP_NS}}}extent")
                img_style = ""
                if extent is not None:
                    cx = extent.get("cx")
                    cy = extent.get("cy")
                    if cx and cy:
                        try:
                            # 1 px = 9525 EMUs (ở 96 DPI)
                            w_px = int(cx) / 9525
                            h_px = int(cy) / 9525
                            img_style = f' style="width: {w_px:.1f}px; height: {h_px:.1f}px;"'
                        except ValueError:
                            pass
                            
                for blip in child.iter(f"{{{A_NS}}}blip"):
                    r_embed = blip.get(f"{{{R_NS}}}embed")
                    if r_embed and r_embed in image_rels:
                        img_name = image_rels[r_embed]
                        src = f"/api/media/{doc_filename}/{img_name}"
                        run_parts.append(f'<img src="{src}" class="doc-img" alt="{img_name}"{img_style} />')
            elif child.tag == f"{{urn:schemas-microsoft-com:vml}}shape":
                imagedata = child.find(f"{{{V_NS}}}imagedata")
                if imagedata is not None:
                    r_id = imagedata.get(f"{{{R_NS}}}id")
                    if r_id and r_id in image_rels:
                        img_name = image_rels[r_id]
                        src = f"/api/media/{doc_filename}/{img_name}"
                        
                        # Trích xuất kích thước của VML shape
                        shape_style = child.get("style", "")
                        img_style = ""
                        if shape_style:
                            styles_vml = []
                            for part in shape_style.split(";"):
                                if "width" in part or "height" in part:
                                    styles_vml.append(part.strip())
                            if styles_vml:
                                img_style = f' style="{"; ".join(styles_vml)}"'
                                
                        run_parts.append(f'<img src="{src}" class="doc-img" alt="{img_name}"{img_style} />')
                        
        content = "".join(run_parts)
        if not content:
            return ""
            
        if is_bold:
            content = f"<strong>{content}</strong>"
        if is_italic:
            content = f"<em>{content}</em>"
        if is_underline:
            content = f"<u>{content}</u>"
            
        return content

    def _replace_paragraph_text(
        self, p_elem: etree._Element, new_text: str, doc
    ) -> None:
        """
        Thay thế toàn bộ text trong một đoạn bằng nội dung mới,
        nhưng giữ lại toàn bộ hình ảnh (w:drawing, v:shape) và định dạng paragraph.
        Loại bỏ tất cả bookmark và các text run cũ để đảm bảo XML hoàn toàn sạch và hợp lệ.
        """
        W = self._W
        import copy

        media_runs = []
        for child in p_elem:
            if child.tag == f"{W}r":
                has_media = False
                for sub in child.iter():
                    if "drawing" in sub.tag or "shape" in sub.tag or "imagedata" in sub.tag:
                        has_media = True
                        break
                if has_media:
                    # Loại bỏ các thẻ w:t bên trong run media nếu có
                    for t_elem in list(child.findall(f"{W}t")):
                        child.remove(t_elem)
                    media_runs.append(copy.deepcopy(child))

        # Lưu lại định dạng paragraph
        pPr = p_elem.find(f"{W}pPr")
        if pPr is not None:
            pPr = copy.deepcopy(pPr)

        # XÓA TRẮNG toàn bộ nội dung paragraph để loại bỏ các thẻ rác, bookmark lỗi, v.v.
        p_elem.clear()

        # 1. Thêm lại định dạng paragraph (phải là thẻ đầu tiên của w:p)
        if pPr is not None:
            p_elem.append(pPr)

        # 2. Thêm lại các run chứa hình ảnh
        for mr in media_runs:
            p_elem.append(mr)

        # 3. Thêm text run thông qua python-docx (đảm bảo sinh XML chuẩn 100%)
        from docx.text.paragraph import Paragraph
        p = Paragraph(p_elem, doc)

        parts = new_text.split("**")
        for idx, part in enumerate(parts):
            if not part:
                continue
            is_bold = (idx % 2 == 1)
            
            run = p.add_run(part)
            if is_bold:
                run.bold = True

    # ================================================================
    #  PUBLIC API
    # ================================================================

    def extract_paragraphs_with_ids(self, file_path: str) -> Dict[str, Any]:
        """
        Trích xuất cấu trúc HTML toàn văn và các đoạn văn kèm ID tuần tự từ file .docx.

        Args:
            file_path: Đường dẫn đến file .docx cần xử lý.

        Returns:
            Dict dạng {"html": "...", "paragraphs": {"p_1": {"text": "...", "html": "..."}, ...}}
        """
        file_path_str = str(file_path)
        path = Path(file_path_str)

        if not path.exists():
            raise FileNotFoundError(f"Không tìm thấy file: {file_path_str}")
        if path.suffix.lower() != ".docx":
            raise ValueError(
                f"File phải có định dạng .docx, nhận được: '{path.suffix}'"
            )

        try:
            doc = Document(file_path_str)
            body = doc.element.body
            
            # Đọc quan hệ hình ảnh
            image_rels = self._load_image_relationships(file_path_str)
            doc_filename = path.name

            # Thực hiện đệ quy toàn bộ body để sinh HTML bố cục chuẩn
            html_out, paragraphs_dict = self._body_to_html_and_paragraphs(body, image_rels, doc_filename)

            logger.info(
                "Đã trích xuất %d đoạn văn từ '%s'", len(paragraphs_dict), file_path_str
            )
            return {
                "html": html_out,
                "paragraphs": paragraphs_dict
            }

        except (FileNotFoundError, ValueError):
            raise
        except Exception as e:
            logger.error(
                "Lỗi khi trích xuất đoạn văn từ '%s': %s", file_path_str, e
            )
            raise RuntimeError(
                f"Không thể trích xuất đoạn văn từ '{file_path_str}': {e}"
            ) from e

    def merge_edited_text(
        self,
        original_file_path: str,
        edited_json: Dict[str, str],
        output_file_path: str,
    ) -> None:
        """
        Gộp nội dung đã chỉnh sửa bởi AI vào file .docx gốc.
        """
        original_path_str = str(original_file_path)
        output_path_str = str(output_file_path)
        path = Path(original_path_str)

        if not path.exists():
            raise FileNotFoundError(
                f"Không tìm thấy file gốc: {original_path_str}"
            )
        if path.suffix.lower() != ".docx":
            raise ValueError(
                f"File phải có định dạng .docx, nhận được: '{path.suffix}'"
            )

        if not edited_json:
            logger.warning(
                "edited_json rỗng — sao chép nguyên file gốc sang '%s'.",
                output_path_str,
            )
            Path(output_path_str).parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(original_path_str, output_path_str)
            return

        try:
            doc = Document(original_path_str)
            body = doc.element.body
            paragraphs = self._get_content_paragraphs(body)

            p_index: Dict[str, etree._Element] = {}
            for idx, (p_elem, _) in enumerate(paragraphs, start=1):
                p_index[f"p_{idx}"] = p_elem

            # Kiểm tra ID không hợp lệ (không phải ID gốc và cũng không phải ID chèn thêm hợp lệ)
            import re
            unknown_ids = []
            for k in edited_json.keys():
                if k in p_index:
                    continue
                # Định dạng chèn thêm: p_{number}_ins_{sub_id}
                if re.match(r"^p_\d+_ins_\w+$", k):
                    target_p = f"p_{k.split('_')[1]}"
                    if target_p in p_index:
                        continue
                unknown_ids.append(k)

            if unknown_ids:
                logger.warning(
                    "Bỏ qua các ID không hợp lệ: %s",
                    sorted(unknown_ids),
                )

            # Khởi tạo bản đồ neo (anchors) để giữ vị trí chèn tiếp theo cho mỗi đoạn gốc
            anchors = {p_id: p_elem for p_id, p_elem in p_index.items()}

            # Hàm sắp xếp các key theo trình tự xuất hiện logic
            def get_sort_key(key: str):
                parts = key.split('_')
                try:
                    p_num = int(parts[1])
                except (IndexError, ValueError):
                    p_num = 99999
                is_ins = 1 if 'ins' in parts else 0
                # Lấy số thứ tự chèn nếu có
                ins_num = 0
                if is_ins and len(parts) > 3:
                    try:
                        ins_num = int(parts[3])
                    except ValueError:
                        pass
                return (p_num, is_ins, ins_num)

            sorted_keys = sorted(edited_json.keys(), key=get_sort_key)
            applied_count = 0
            W = self._W

            for key in sorted_keys:
                if key in unknown_ids:
                    continue

                new_text = edited_json[key]

                if "ins" in key:
                    # Logic chèn thêm đoạn văn mới
                    parts = key.split('_')
                    target_id = f"p_{parts[1]}"
                    if target_id not in anchors:
                        continue
                    
                    anchor_elem = anchors[target_id]
                    parent = anchor_elem.getparent()
                    if parent is None:
                        continue

                    from docx.oxml.shared import OxmlElement
                    from docx.text.paragraph import Paragraph

                    # Tạo đoạn văn mới w:p
                    new_p = OxmlElement('w:p')
                    
                    # Sao chép định dạng paragraph w:pPr từ đoạn neo gốc
                    orig_elem = p_index[target_id]
                    orig_p_pr = orig_elem.find(f"{W}pPr")
                    if orig_p_pr is not None:
                        new_p.append(copy.deepcopy(orig_p_pr))

                    # Chèn vào XML sau phần tử neo hiện tại
                    idx_pos = parent.index(anchor_elem)
                    parent.insert(idx_pos + 1, new_p)

                    p = Paragraph(new_p, doc)

                    # Phân tích và tạo các run thường & run in đậm dựa trên cú pháp **
                    text_parts = new_text.split("**")
                    for idx, part in enumerate(text_parts):
                        if not part:
                            continue
                        is_bold = (idx % 2 == 1)
                        run = p.add_run(part)
                        
                        if is_bold:
                            run.bold = True

                    # Cập nhật neo để lượt chèn tiếp theo nằm sau đoạn mới này
                    anchors[target_id] = new_p
                    applied_count += 1
                else:
                    # Logic sửa đổi hoặc xóa đoạn văn gốc
                    if key not in p_index:
                        continue
                    p_elem = p_index[key]

                    if new_text == "":
                        # Thay vì xóa hoàn toàn w:p (gây lỗi cấu trúc XML nếu là đoạn cuối cùng của Table Cell hoặc Body), 
                        # ta chỉ xóa nội dung để biến nó thành đoạn trống an toàn.
                        self._replace_paragraph_text(p_elem, "", doc)
                        
                        # Xóa khỏi danh sách neo để không thể chèn sau nó nữa
                        if key in anchors:
                            del anchors[key]
                        applied_count += 1
                    else:
                        # Thay thế văn bản
                        self._replace_paragraph_text(p_elem, new_text, doc)
                        applied_count += 1

            Path(output_path_str).parent.mkdir(parents=True, exist_ok=True)
            doc.save(output_path_str)

            logger.info(
                "Đã gộp %d/%d chỉnh sửa → '%s'",
                applied_count,
                len(edited_json),
                output_path_str,
            )

        except (FileNotFoundError, ValueError):
            raise
        except Exception as e:
            logger.error("Lỗi khi gộp chỉnh sửa: %s", e)
            out = Path(output_path_str)
            if out.exists():
                try:
                    out.unlink()
                    logger.info("Đã xóa file đầu ra lỗi: '%s'", output_path_str)
                except OSError:
                    pass
            raise RuntimeError(f"Không thể gộp chỉnh sửa: {e}") from e
