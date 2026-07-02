from docx import Document
from docx.oxml.shared import OxmlElement
from docx.text.paragraph import Paragraph
import copy

doc = Document()
doc.add_paragraph('Test Paragraph')
body = doc.element.body
anchor = body[0]
new_p = OxmlElement('w:p')
body.insert(body.index(anchor) + 1, new_p)

# Create a mock first_r_pr with w:sz
first_r_pr = OxmlElement('w:rPr')
sz = OxmlElement('w:sz')
sz.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '24')
first_r_pr.append(sz)

p = Paragraph(new_p, doc)
run = p.add_run('World')

# Set style
# run._r usually has no rPr at creation if no style is given
rPr = run._r.get_or_add_rPr()
run._r.remove(rPr)
run._r.append(copy.deepcopy(first_r_pr))

# Now set bold
run.bold = True

print(run._r.xml)
