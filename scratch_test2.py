from docx import Document
from docx.text.paragraph import Paragraph

doc = Document()
doc.add_paragraph('Test Paragraph')

body = doc.element.body
p_elem = body[0]

try:
    p = Paragraph(p_elem, doc)
    run = p.add_run('World')
    run.bold = True
    print('Success:', p_elem.tag)
except Exception as e:
    print('Error:', e)
