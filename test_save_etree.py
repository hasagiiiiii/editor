from docx import Document
from lxml import etree
import copy

doc = Document()
p = doc.add_paragraph()
run = p.add_run('part')

first_r_pr = etree.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
sz = etree.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz')
sz.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '24')
first_r_pr.append(sz)

rPr = run._r.get_or_add_rPr()
run._r.remove(rPr)
run._r.insert(0, copy.deepcopy(first_r_pr))
run.bold = True

doc.save('d:/editor/test_corrupt.docx')
print("Saved")
